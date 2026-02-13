import os
from docx import Document
from openpyxl import Workbook

# Set the directory of your project
project_dir = r"D:\Machine_Learning_Projects\7. GARGI – Guided AI for Real-world Grammar & Interaction"

# Output Word document file for project details
output_file_docx = "project_details_API.docx"
# Output Excel file for file paths
output_file_excel = "project_files_API.xlsx"

# Folder and file types to ignore
ignore_folders = (".git", "venv", "__pycache__", "data", "output", "logs", "models")
ignore_files = ("LICENSE", "README.md", "git", ".gitignore", ".docx", ".pyc", "__init__.py",
                 ".xlsx", "test1_Android.py", "test2_API.py", )

# Function to clean up file contents (remove non-XML characters)``
def clean_content(content):
    # Remove non-printable characters or any control characters
    return ''.join([char if char.isprintable() else ' ' for char in content])

# Function to get file contents 
def get_file_contents(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = clean_content(file.read())
            return content  # Limit content to first 1000 characters
    except Exception as e:
        return f"Could not read file {file_path}. Error: {e}"

# Function to collect file paths and content, then save them to Word and Excel
def collect_project_files_and_save(directory, word_filename, excel_filename):
    # Create Word document for file details
    doc = Document()
    doc.add_heading('Project Details', 0)

    # Create Excel workbook for file paths
    wb = Workbook()
    ws = wb.active
    ws.title = "File Paths"

    # Add header to Excel file
    ws.append(["File Path"])

    # Walk through the directory to collect relevant file paths and content
    for root, dirs, files in os.walk(directory):
        # Skip ignored folders
        dirs[:] = [d for d in dirs if not any(ignored in d for ignored in ignore_folders)]

        for file_name in files:
            file_path = os.path.join(root, file_name)

            # Skip ignored files (like LICENSE, README.md, etc.)
            if any(file_name.endswith(ext) for ext in ignore_files):
                continue

            # Only include .py and .jsonl files
            if file_name.endswith(".py") or file_name.endswith(".jsonl"):
                # Add file path to Excel
                ws.append([file_path])

                # Add file path and content to Word document
                doc.add_paragraph(f"File Path: {file_path}")
                content = get_file_contents(file_path)
                doc.add_paragraph(f"Content: {content}")
                doc.add_paragraph("=" * 40)  # Separator line for readability

    # Save Word document and Excel file
    doc.save(word_filename)
    wb.save(excel_filename)

    print(f"Project details saved to {word_filename}")
    print(f"File paths saved to {excel_filename}")

# Run the function to collect files and save them
collect_project_files_and_save(project_dir, output_file_docx, output_file_excel)
