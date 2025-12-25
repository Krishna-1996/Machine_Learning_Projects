import os
from docx import Document
from openpyxl import Workbook

# Set the directory of your project
project_dir = r"D:\Android\Andriod_Data\AndroidStudioProjects\GargiAndroid"

# Output Word document file
output_file_docx = "project_details.docx"
# Output Excel file for folder paths
output_file_excel = "folder_paths.xlsx"

# Folder to ignore
ignore_folder = "build"

# Function to clean up file contents (remove non-XML characters)
def clean_content(content):
    # Remove non-printable characters or any control characters
    return ''.join([char if char.isprintable() else ' ' for char in content])

# Function to get file contents
def get_file_contents(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return clean_content(file.read())
    except Exception as e:
        return f"Could not read file {file_path}. Error: {e}"

# Function to collect folder paths and save them to Excel
def save_folder_paths_to_excel(directory, excel_filename):
    wb = Workbook()
    ws = wb.active
    ws.title = "Folder Paths"

    # Add the header
    ws.append(["Folder Path"])

    # Walk through the directory and add folder paths to the Excel file
    for root, dirs, files in os.walk(directory):
        # Skip any directory named "build" or any directory starting with a dot
        dirs[:] = [d for d in dirs if ignore_folder not in d and not d.startswith('.')]

        # Add folder path to Excel
        ws.append([root])

    # Save the workbook
    wb.save(excel_filename)
    print(f"Folder paths saved to {excel_filename}")

# Create a new Word document for project details
doc = Document()
doc.add_heading('Project Details', 0)

# Collect details from the project folder and add them to the Word document
def collect_project_details(directory, doc):
    for root, dirs, files in os.walk(directory):
        # Skip any directory named "build" or any directory starting with a dot
        dirs[:] = [d for d in dirs if ignore_folder not in d and not d.startswith('.')]

        for file_name in files:
            file_path = os.path.join(root, file_name)

            # Add File Path
            doc.add_paragraph(f"File Path: {file_path}")
            
            # Add File Content (limit to first 1000 characters)
            content = get_file_contents(file_path)
            doc.add_paragraph(f"File Content: {content}")  # Limiting to first 1000 characters to avoid too much content
            doc.add_paragraph("=" * 40)  # Separator line for readability

# Run the function to collect details and add them to the Word document
collect_project_details(project_dir, doc)

# Save the Word document
doc.save(output_file_docx)

# Save folder paths to Excel
save_folder_paths_to_excel(project_dir, output_file_excel)

print(f"Project details saved to {output_file_docx}")
